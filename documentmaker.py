import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def maker(_sn,_nm,_sap,_rn,_fn):
    doc = docx.Document()
    heading = doc.add_heading(_sn,0)
    heading.style.font.name = "Times New Roman"
    heading.style.font.size = Pt(30)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _name = doc.add_paragraph("Name: "+_nm)
    _name.style.font.name = "Times New Roman"
    _name.style.font.size = Pt(14)

    _sid = doc.add_paragraph("SAP ID: "+_sap)
    _sid.style.font.name = "Times New Roman"
    _sid.style.font.size = Pt(14)

    _rn = doc.add_paragraph("Roll No: "+_rn)
    _rn.style.font.name = "Times New Roman"
    _rn.style.font.size = Pt(14)

    return doc.save('{}.docx'.format(_fn))

